"""Genera diagramas de actividad en Word como tablas de flujo."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54)
    s.right_margin = Cm(2.54)

h = doc.add_heading('5.3.3 Diagramas de Actividad', level=2)
for r in h.runs:
    r.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph(
    'Los diagramas de actividad modelan el flujo de trabajo de los procesos mas representativos '
    'del sistema SIGAE. Cada diagrama muestra la secuencia de acciones, las decisiones que bifurcan '
    'el flujo y los estados finales, aplicando la notacion UML estandar.'
)

def set_cell(cell, text, bold=False, size=9, bg=None, color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    if bg:
        from docx.oxml.ns import qn
        tc = cell._element.get_or_add_tcPr()
        sh = tc.makeelement(qn('w:shd'), {qn('w:fill'): bg, qn('w:val'): 'clear'})
        tc.append(sh)

def add_activity(title, steps):
    """
    steps: list of (tipo, texto)
    tipo: 'inicio', 'fin', 'accion', 'decision', 'si', 'no', 'atomic', 'error', 'exito'
    """
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 48, 135)

    t = doc.add_table(rows=len(steps), cols=3)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (tipo, texto) in enumerate(steps):
        row = t.rows[i]
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(1.5)
        row.cells[2].width = Cm(13.5)

        if tipo == 'inicio':
            set_cell(row.cells[0], '\u25CF', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0,0,0))
            set_cell(row.cells[1], '\u2193', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(row.cells[2], texto, bold=True, size=9, bg='003087', color=RGBColor(255,255,255))
        elif tipo == 'fin':
            set_cell(row.cells[0], '\u25C9', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0,0,0))
            set_cell(row.cells[1], '', size=9)
            set_cell(row.cells[2], texto, bold=True, size=9, bg='333333', color=RGBColor(255,255,255))
        elif tipo == 'accion':
            set_cell(row.cells[0], '\u25A2', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0,48,135))
            set_cell(row.cells[1], '\u2193', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(row.cells[2], texto, size=9, bg='E8F4FD')
        elif tipo == 'decision':
            set_cell(row.cells[0], '\u25C7', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(243,156,18))
            set_cell(row.cells[1], '\u2193', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(row.cells[2], texto, bold=True, size=9, bg='FFF3CD')
        elif tipo == 'si':
            set_cell(row.cells[0], '', size=9)
            set_cell(row.cells[1], 'Si \u2192', size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(39,174,96))
            set_cell(row.cells[2], texto, size=9, bg='D4EDDA')
        elif tipo == 'no':
            set_cell(row.cells[0], '', size=9)
            set_cell(row.cells[1], 'No \u2192', size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(192,57,43))
            set_cell(row.cells[2], texto, size=9, bg='FEE2E2')
        elif tipo == 'atomic':
            set_cell(row.cells[0], '\U0001F512', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(row.cells[1], '\u2193', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(row.cells[2], texto, bold=True, size=9, bg='D4EDDA')
        elif tipo == 'error':
            set_cell(row.cells[0], '\u2716', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(192,57,43))
            set_cell(row.cells[1], '', size=9)
            set_cell(row.cells[2], texto, size=9, bg='FEE2E2')
        elif tipo == 'exito':
            set_cell(row.cells[0], '\u2714', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(39,174,96))
            set_cell(row.cells[1], '\u2193', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(row.cells[2], texto, size=9, bg='D4EDDA')

    doc.add_paragraph()


# ================================================================
# DA-01: Proceso de Gestion de Activos
# ================================================================
add_activity('DA-01: Proceso de Gestion de Activos (CU-02)', [
    ('inicio', 'INICIO'),
    ('accion', 'El usuario accede al modulo "Activos" desde el menu lateral del sistema'),
    ('decision', 'El usuario tiene rol OPERADOR?'),
    ('si', 'El sistema filtra automaticamente los activos mostrando solo los de la sede asignada al operador (sede_id)'),
    ('no', 'El sistema muestra todos los activos de todas las sedes (vista de Administrador)'),
    ('accion', 'El sistema presenta la lista de activos con barra de busqueda (nombre/serial/tipo) y filtro por estado (Disponible/Asignado/Reparacion)'),
    ('decision', 'Que accion selecciona el usuario?'),
    ('si', '[CREAR] El sistema despliega formulario vacio con campos: nombre, descripcion, tipo, serial, valor, fecha compra, sede y departamento'),
    ('accion', 'El usuario completa los campos obligatorios y hace clic en "Guardar"'),
    ('decision', 'Los datos son validos? (serial unico, valor > 0, campos completos)'),
    ('si', 'El sistema crea el activo con estado DISPONIBLE y registra la operacion en el log de actividad'),
    ('no', 'El sistema muestra el mensaje de error correspondiente: "Serial ya existe", "Valor debe ser mayor a 0" o "Complete todos los campos"'),
    ('exito', 'Notificacion: "Activo creado exitosamente". Se actualiza la tabla'),
    ('accion', '[EDITAR] El usuario hace clic en "Editar" de un activo existente. El sistema carga los datos en el formulario incluyendo el campo de estado'),
    ('accion', '[ELIMINAR] El usuario hace clic en "Eliminar" de un activo'),
    ('decision', 'El activo tiene asignaciones activas?'),
    ('si', 'El sistema muestra: "No se puede eliminar: esta asignado" y bloquea la operacion'),
    ('no', 'El sistema solicita confirmacion, elimina el activo y registra en el log'),
    ('fin', 'FIN'),
])

# ================================================================
# DA-02: Proceso de Asignacion y Devolucion
# ================================================================
add_activity('DA-02: Proceso de Asignacion y Devolucion de Activos (CU-04)', [
    ('inicio', 'INICIO - ASIGNACION'),
    ('accion', 'El sistema carga la lista de activos con estado DISPONIBLE y la lista de empleados (filtrados por sede si es OPERADOR)'),
    ('decision', 'Hay activos disponibles?'),
    ('no', 'El formulario se deshabilita mostrando "Sin activos disponibles". El usuario solo puede consultar el historial'),
    ('si', 'El usuario selecciona un activo del dropdown, un empleado y opcionalmente escribe observaciones'),
    ('accion', 'El usuario hace clic en "Registrar"'),
    ('decision', 'El activo seleccionado sigue en estado DISPONIBLE? (verificacion en backend)'),
    ('no', 'El sistema muestra: "Activo no disponible" (otro usuario pudo asignarlo simultaneamente)'),
    ('atomic', 'TRANSACCION ATOMICA: 1) Crear registro de asignacion con fecha_asignacion = ahora. 2) Cambiar estado del activo de DISPONIBLE a ASIGNADO'),
    ('accion', 'El sistema registra la operacion en el log de actividad'),
    ('exito', 'Notificacion: "Asignacion registrada". Se actualiza la tabla y el dropdown de disponibles'),
    ('inicio', 'INICIO - DEVOLUCION'),
    ('accion', 'El usuario localiza una asignacion con estado "Activa" en la tabla de historial'),
    ('accion', 'El usuario hace clic en el boton "Devolver"'),
    ('decision', 'El usuario confirma la devolucion en el dialogo?'),
    ('no', 'Se cancela la operacion. No se realizan cambios'),
    ('atomic', 'TRANSACCION ATOMICA: 1) Registrar fecha_devolucion = ahora en la asignacion. 2) Cambiar estado del activo de ASIGNADO a DISPONIBLE'),
    ('accion', 'El sistema registra la operacion en el log de actividad'),
    ('exito', 'Notificacion: "Devolucion registrada". El estado cambia a "Devuelta" en la tabla'),
    ('fin', 'FIN'),
])

# ================================================================
# DA-03: Proceso de Autenticacion
# ================================================================
add_activity('DA-03: Proceso de Autenticacion (CU-01)', [
    ('inicio', 'INICIO'),
    ('accion', 'El usuario accede a la pantalla de login del sistema'),
    ('accion', 'El usuario ingresa su nombre de usuario y contrasena en los campos del formulario'),
    ('accion', 'El usuario hace clic en el boton "Ingresar"'),
    ('accion', 'El frontend envia las credenciales al backend mediante POST /api/token/'),
    ('decision', 'Las credenciales son validas? (username existe y hash de contrasena coincide)'),
    ('no', 'El sistema muestra: "Usuario o contrasena incorrectos" sin revelar cual es incorrecto. El usuario puede intentar nuevamente'),
    ('si', 'El backend genera tokens JWT: access (expiracion corta) y refresh (expiracion larga)'),
    ('accion', 'El frontend almacena los tokens en localStorage y carga la sesion en AuthContext con los datos del usuario (id, username, rol, sede_id)'),
    ('decision', 'Cual es el rol del usuario autenticado?'),
    ('si', '[ADMIN] El sistema muestra el menu completo incluyendo la seccion Administracion (Usuarios, Actividad)'),
    ('no', '[OPERADOR] El sistema muestra el menu sin la seccion Administracion. Los datos se filtran por sede_id'),
    ('exito', 'Redireccion exitosa a /dashboard con el sidebar adaptado segun el rol'),
    ('fin', 'FIN'),
])

# ================================================================
# DA-04: Proceso de Gestion de Sedes
# ================================================================
add_activity('DA-04: Proceso de Gestion de Sedes (CU-05) - Solo Admin', [
    ('inicio', 'INICIO'),
    ('decision', 'El usuario tiene rol ADMIN?'),
    ('no', 'El modulo de Sedes no es visible en el menu. Acceso denegado'),
    ('si', 'El sistema muestra la tabla de sedes con: nombre, ciudad, direccion, telefono, cantidad de departamentos y activos'),
    ('accion', 'El administrador hace clic en "+ Nueva Sede"'),
    ('accion', 'El sistema despliega formulario: nombre (obligatorio), ciudad (obligatorio), direccion (obligatorio), telefono (opcional)'),
    ('accion', 'El administrador completa los campos y hace clic en "Guardar"'),
    ('decision', 'Los campos obligatorios estan completos?'),
    ('no', 'El sistema muestra: "Complete los campos obligatorios"'),
    ('si', 'El sistema crea la sede, registra en el log de actividad y actualiza la tabla'),
    ('exito', 'Notificacion: "Sede guardada exitosamente"'),
    ('accion', '[ELIMINAR] El administrador hace clic en "Eliminar" de una sede'),
    ('decision', 'La sede tiene activos, empleados o departamentos asociados?'),
    ('si', 'El sistema muestra: "No se puede eliminar: la sede tiene [activos/empleados/departamentos] asociados"'),
    ('no', 'El sistema solicita confirmacion, elimina la sede y registra en el log'),
    ('fin', 'FIN'),
])

# ================================================================
# DA-05: Proceso de Generacion de Reportes
# ================================================================
add_activity('DA-05: Proceso de Generacion de Reportes (CU-06)', [
    ('inicio', 'INICIO'),
    ('accion', 'El usuario accede al modulo "Reportes" desde el menu lateral'),
    ('accion', 'El sistema consulta todos los activos (filtrados por sede si es OPERADOR) y calcula estadisticas'),
    ('accion', 'El sistema muestra 4 tarjetas de resumen: valor total del inventario, total de activos, asignados y en reparacion'),
    ('accion', 'El sistema muestra tabs de navegacion: Disponibles, Asignados, Reparacion'),
    ('accion', 'El usuario selecciona un tab para ver el reporte correspondiente'),
    ('accion', 'El sistema muestra la tabla con los activos filtrados por el estado seleccionado'),
    ('decision', 'El usuario desea exportar los datos?'),
    ('no', 'El usuario consulta los datos en pantalla. Puede cambiar de tab para ver otros reportes'),
    ('si', 'El usuario hace clic en "Exportar CSV"'),
    ('accion', 'El sistema genera un archivo CSV con codificacion UTF-8 y BOM para compatibilidad con Excel'),
    ('exito', 'El archivo activos_reporte.csv se descarga automaticamente en el navegador'),
    ('fin', 'FIN'),
])

# ================================================================
# DA-06: Proceso de Gestion de Usuarios
# ================================================================
add_activity('DA-06: Proceso de Gestion de Usuarios (CU-07) - Solo Admin', [
    ('inicio', 'INICIO'),
    ('decision', 'El usuario tiene rol ADMIN?'),
    ('no', 'La seccion Administracion no es visible en el menu. Acceso denegado'),
    ('si', 'El sistema muestra el formulario de creacion y la tabla de usuarios existentes'),
    ('accion', 'El administrador completa: usuario, email, contrasena, nombre, apellido, rol (Operador/Administrador) y sede'),
    ('accion', 'El administrador hace clic en "Crear Usuario"'),
    ('decision', 'Los datos son validos? (username unico, contrasena >= 6 caracteres, email valido)'),
    ('no', 'El sistema muestra el error correspondiente: "Usuario ya existe" o "Contrasena minimo 6 caracteres"'),
    ('si', 'El sistema crea el usuario con la contrasena cifrada (PBKDF2-SHA256) y registra en el log'),
    ('exito', 'Notificacion: "Usuario creado: [username]". Se actualiza la tabla'),
    ('accion', '[ELIMINAR] El administrador hace clic en "Eliminar" de un usuario'),
    ('decision', 'El usuario a eliminar es el mismo administrador logueado?'),
    ('si', 'El sistema muestra: "No puede eliminar su propia cuenta" y bloquea la operacion'),
    ('no', 'El sistema solicita confirmacion, elimina el usuario y registra en el log'),
    ('fin', 'FIN'),
])

# ================================================================
# DA-07: Proceso de Gestion de Empleados
# ================================================================
add_activity('DA-07: Proceso de Gestion de Empleados (CU-03)', [
    ('inicio', 'INICIO'),
    ('accion', 'El usuario accede al modulo "Empleados" desde el menu lateral'),
    ('accion', 'El sistema muestra la lista de empleados (filtrada por sede si es OPERADOR) con barra de busqueda'),
    ('accion', 'El usuario hace clic en "+ Nuevo Empleado"'),
    ('accion', 'El sistema despliega formulario: nombre, apellido, cedula, email, sede, departamento y cargo'),
    ('accion', 'El usuario completa los campos y hace clic en "Guardar"'),
    ('decision', 'Los datos son validos? (cedula unica, campos obligatorios completos)'),
    ('no', 'El sistema muestra: "Cedula ya existe" o "Complete los campos obligatorios"'),
    ('si', 'El sistema crea el empleado con estado activo=true y fecha_ingreso automatica. Registra en log'),
    ('exito', 'Notificacion: "Empleado creado exitosamente". Se actualiza la tabla'),
    ('accion', '[DAR DE BAJA] El usuario hace clic en "Baja" de un empleado activo'),
    ('accion', 'El sistema registra fecha_salida y marca activo=false. El empleado aparece atenuado en la tabla'),
    ('accion', '[REACTIVAR] El usuario hace clic en "Reactivar" de un empleado dado de baja'),
    ('accion', 'El sistema limpia fecha_salida y marca activo=true'),
    ('accion', '[ELIMINAR] El usuario hace clic en "Eliminar"'),
    ('decision', 'El empleado tiene activos asignados actualmente?'),
    ('si', 'El sistema muestra: "No se puede eliminar: tiene activos asignados"'),
    ('no', 'El sistema solicita confirmacion, elimina el empleado y registra en el log'),
    ('fin', 'FIN'),
])

output = 'documentacion/Diagramas_Actividad_SIGAE.docx'
doc.save(output)
print(f'Documento generado: {output}')
print(f'Tablas: {len(doc.tables)}')
