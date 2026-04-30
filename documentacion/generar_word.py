"""Convierte Informe_Final_SIGAE.md a Word (.docx) con formato academico."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Estilos globales
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2.54)


def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


# Leer el markdown
with open('documentacion/Informe_Final_SIGAE.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_table = False
table_headers = []
table_rows = []
in_code = False
code_lines = []

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Code blocks
    if line.startswith('```'):
        if in_code:
            # End code block
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            code_lines = []
            in_code = False
        else:
            in_code = True
            code_lines = []
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # Tables
    if '|' in line and not line.startswith('```'):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells and all(c.strip() == '' or set(c.strip()) <= {'-', ':', ' '} for c in cells):
            # Separator row, skip
            i += 1
            continue
        if not in_table:
            in_table = True
            table_headers = cells
            table_rows = []
        else:
            table_rows.append(cells)
        # Check if next line is still table
        if i + 1 < len(lines) and '|' in lines[i+1]:
            i += 1
            continue
        else:
            # End of table, render it
            if table_headers and table_rows:
                # Pad rows to match headers
                for ri in range(len(table_rows)):
                    while len(table_rows[ri]) < len(table_headers):
                        table_rows[ri].append('')
                add_table(table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
            i += 1
            continue

    # If we were in a table but line doesn't have |, flush
    if in_table:
        if table_headers and table_rows:
            for ri in range(len(table_rows)):
                while len(table_rows[ri]) < len(table_headers):
                    table_rows[ri].append('')
            add_table(table_headers, table_rows)
        in_table = False
        table_headers = []
        table_rows = []

    # Headings
    if line.startswith('# ') and not line.startswith('##'):
        text = line[2:].strip()
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0, 0, 0)
        i += 1
        continue

    if line.startswith('## '):
        text = line[3:].strip()
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0, 0, 0)
        i += 1
        continue

    if line.startswith('### '):
        text = line[4:].strip()
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0, 0, 0)
        i += 1
        continue

    if line.startswith('#### '):
        text = line[5:].strip()
        h = doc.add_heading(text, level=4)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0, 0, 0)
        i += 1
        continue

    # Horizontal rules / page breaks
    if line.strip() == '---':
        doc.add_page_break()
        i += 1
        continue

    # Bullet points
    if line.startswith('- '):
        text = line[2:].strip()
        # Remove markdown bold
        text = text.replace('**', '')
        doc.add_paragraph(text, style='List Bullet')
        i += 1
        continue

    # Numbered lists
    m = re.match(r'^(\d+)\.\s+(.+)', line)
    if m:
        text = m.group(2).replace('**', '')
        doc.add_paragraph(text, style='List Number')
        i += 1
        continue

    # Empty lines
    if line.strip() == '':
        i += 1
        continue

    # Regular paragraphs
    text = line.strip()
    # Remove markdown formatting
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'>\s*', '', text)

    if text:
        doc.add_paragraph(text)

    i += 1

# Guardar
output = 'documentacion/Informe_SIGAE_Final.docx'
doc.save(output)
print(f'Documento generado: {output}')
print(f'Parrafos: {len(doc.paragraphs)}')
print(f'Tablas: {len(doc.tables)}')
