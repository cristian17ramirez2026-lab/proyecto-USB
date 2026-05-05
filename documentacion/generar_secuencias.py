"""Genera diagramas de secuencia en Word como tablas formateadas."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54)
    s.right_margin = Cm(2.54)

h = doc.add_heading('5.3.2 Diagramas de Secuencia', level=2)
for r in h.runs:
    r.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph(
    'Los diagramas de secuencia representan la interaccion temporal entre los actores y los '
    'componentes del sistema para cada caso de uso. Se presenta un diagrama por cada caso de uso, '
    'mostrando el flujo de mensajes entre el usuario, el frontend (React), el backend (Django REST API) '
    'y la base de datos (SQLite). Los mensajes sincronos se representan con flecha solida (\u2192), '
    'las respuestas con flecha punteada (\u21E0), y las operaciones internas con auto-mensaje (\u21BB).'
)

def set_cell(cell, text, bold=False, size=9, bg=None, color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    if bg:
        from docx.oxml.ns import qn
        tc = cell._element.get_or_add_tcPr()
        sh = tc.makeelement(qn('w:shd'), {qn('w:fill'): bg, qn('w:val'): 'clear'})
        tc.append(sh)

def add_seq_diagram(title, ref, participants, messages, notes=None):
    """
    participants: list of str
    messages: list of (num, from_idx, to_idx, arrow, text, is_fragment)
      arrow: '->' sincrono, '<--' respuesta, 'self' auto-mensaje
      is_fragment: None or 'atomic' or 'alt'
    notes: list of str (optional)
    """
    # Titulo
    p = doc.add_paragraph()
    run = p.add_run(f'{title} ({ref})')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 48, 135)

    # Tabla de participantes
    pt = doc.add_table(rows=1, cols=len(participants))
    pt.style = 'Table Grid'
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, name in enumerate(participants):
        set_cell(pt.rows[0].cells[i], name, bold=True, size=9, bg='003087',
                 color=RGBColor(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

    # Tabla de mensajes
    mt = doc.add_table(rows=1 + len(messages), cols=5)
    mt.style = 'Table Grid'
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ['#', 'Origen', 'Dir.', 'Destino', 'Mensaje / Accion']
    widths = [Cm(1), Cm(3), Cm(1.2), Cm(3), Cm(9.3)]
    for i, h_text in enumerate(headers):
        c = mt.rows[0].cells[i]
        c.width = widths[i]
        set_cell(c, h_text, bold=True, size=8, bg='E8EEF5', align=WD_ALIGN_PARAGRAPH.CENTER)

    # Rows
    for idx, msg in enumerate(messages):
        num, fr, to, arrow, text, frag = msg
        row = mt.rows[idx + 1]

        # Numero
        set_cell(row.cells[0], str(num), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Determinar nombres
        if arrow == 'self':
            fr_name = participants[fr]
            to_name = participants[fr]
            arrow_sym = '\u21BB'
        elif arrow == '->':
            fr_name = participants[fr]
            to_name = participants[to]
            arrow_sym = '\u2192'
        else:  # '<--'
            fr_name = participants[to]
            to_name = participants[fr]
            arrow_sym = '\u21E0'

        set_cell(row.cells[1], fr_name, size=9)
        set_cell(row.cells[2], arrow_sym, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0, 48, 135))
        set_cell(row.cells[3], to_name, size=9)

        # Texto con fondo verde si es atomic
        bg = None
        if frag == 'atomic':
            bg = 'D4EDDA'
            text = '\U0001F512 ' + text
        set_cell(row.cells[4], text, size=9, bg=bg)

    # Notas
    if notes:
        for note in notes:
            p = doc.add_paragraph()
            run = p.add_run(note)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.italic = True

    doc.add_paragraph()  # Espacio


P = ['Usuario', 'React Frontend', 'Django REST API', 'SQLite BD']
PL = ['Admin', 'React Frontend', 'LocalStorage DB']
PU = ['Usuario', 'React Frontend', 'LocalStorage DB']

# SD-01: Iniciar Sesion
add_seq_diagram('SD-01: Iniciar Sesion', 'CU-01', P, [
    (1, 0, 1, '->', 'Ingresa usuario y contrasena, clic en "Ingresar"', None),
    (2, 1, 2, '->', 'POST /api/token/ {username, password}', None),
    (3, 2, 3, '->', 'SELECT usuario WHERE username=? + verificar hash PBKDF2', None),
    (4, 3, 2, '<--', 'usuario {id, username, email, rol} si credenciales validas', None),
    (5, 2, 2, 'self', 'Generar tokens JWT (access + refresh)', None),
    (6, 2, 1, '<--', '200 OK {access, refresh, user: {id, username, rol}}', None),
    (7, 1, 1, 'self', 'localStorage.setItem(tokens) + AuthContext.setUser(session)', None),
    (8, 1, 0, '<--', 'Redirect a /dashboard. Sidebar adaptado segun rol', None),
], ['Flujo alterno: credenciales invalidas -> 401 Unauthorized -> "Usuario o contrasena incorrectos"'])

# SD-02: Crear Activo
add_seq_diagram('SD-02: Crear Activo', 'CU-02', P, [
    (1, 0, 1, '->', 'Completa formulario (nombre, tipo, serial, valor, fecha, sede, depto)', None),
    (2, 1, 2, '->', 'POST /api/activos/ {data} + Header Authorization: Bearer JWT', None),
    (3, 2, 2, 'self', 'Verificar token JWT valido y no expirado', None),
    (4, 2, 3, '->', 'SELECT * FROM activo WHERE serial=? (verificar unicidad)', None),
    (5, 3, 2, '<--', 'Resultado vacio: serial disponible', None),
    (6, 2, 2, 'self', 'ActivoSerializer.validate(): valor > 0, estado valido, campos completos', None),
    (7, 2, 3, '->', 'INSERT INTO activo (...) VALUES (...) con estado=DISPONIBLE', None),
    (8, 3, 2, '<--', 'Activo creado {id}', None),
    (9, 2, 1, '<--', '201 Created {activo_data completa}', None),
    (10, 1, 1, 'self', 'setState() actualizar tabla + DB.addLog("Activo creado")', None),
    (11, 1, 0, '<--', 'Notificacion: "Activo creado exitosamente"', None),
], ['Flujo alterno: serial duplicado -> 400 Bad Request -> "Ya existe un activo con este serial"'])

# SD-03: Crear Empleado
add_seq_diagram('SD-03: Crear Empleado', 'CU-03', P, [
    (1, 0, 1, '->', 'Completa formulario (nombre, apellido, cedula, email, sede, depto, cargo)', None),
    (2, 1, 2, '->', 'POST /api/empleados/ {data} + JWT', None),
    (3, 2, 3, '->', 'SELECT * FROM empleado WHERE cedula=? (verificar unicidad)', None),
    (4, 3, 2, '<--', 'Resultado vacio: cedula disponible', None),
    (5, 2, 2, 'self', 'EmpleadoSerializer.validate(): campos obligatorios completos', None),
    (6, 2, 3, '->', 'INSERT INTO empleado (...) VALUES (...)', None),
    (7, 3, 2, '<--', 'Empleado creado {id}', None),
    (8, 2, 1, '<--', '201 Created {empleado_data}', None),
    (9, 1, 1, 'self', 'setState() + DB.addLog("Empleado creado")', None),
    (10, 1, 0, '<--', 'Notificacion: "Empleado creado exitosamente"', None),
], ['Flujo alterno: cedula duplicada -> 400 Bad Request -> "Ya existe un empleado con esta cedula"'])

# SD-04: Asignar Activo
add_seq_diagram('SD-04: Asignar Activo a Empleado', 'CU-04', P, [
    (1, 0, 1, '->', 'Selecciona activo disponible + empleado + observaciones. Clic "Registrar"', None),
    (2, 1, 2, '->', 'POST /api/asignaciones/ {activo_id, empleado_id, observaciones} + JWT', None),
    (3, 2, 2, 'self', 'Verificar token JWT', None),
    (4, 2, 3, '->', 'SELECT estado FROM activo WHERE id=activo_id', None),
    (5, 3, 2, '<--', 'estado = DISPONIBLE (validacion exitosa)', None),
    (6, 2, 3, '->', '@transaction.atomic: INSERT INTO asignacion (fecha=NOW)', 'atomic'),
    (7, 2, 3, '->', '@transaction.atomic: UPDATE activo SET estado=ASIGNADO', 'atomic'),
    (8, 3, 2, '<--', 'COMMIT exitoso (ambas operaciones aplicadas)', 'atomic'),
    (9, 2, 1, '<--', '201 Created {asignacion con activo y empleado serializados}', None),
    (10, 1, 1, 'self', 'Actualizar tabla de asignaciones + refrescar dropdown de disponibles', None),
    (11, 1, 0, '<--', 'Notificacion: "Asignacion registrada"', None),
], ['Nota: Los pasos 6-8 se ejecutan como transaccion atomica. Si alguno falla, ninguno se aplica.',
    'Flujo alterno: activo ya asignado -> 400 "El activo ya esta asignado a otro empleado"'])

# SD-05: Devolucion
add_seq_diagram('SD-05: Registrar Devolucion', 'CU-04', P, [
    (1, 0, 1, '->', 'Localiza asignacion activa en la tabla. Clic en "Devolver"', None),
    (2, 1, 0, '<--', 'Dialogo de confirmacion: "Confirmar devolucion?"', None),
    (3, 0, 1, '->', 'Confirma la devolucion', None),
    (4, 1, 2, '->', 'POST /api/asignaciones/{id}/devolver/ + JWT', None),
    (5, 2, 3, '->', 'SELECT asignacion WHERE id=? (verificar fecha_devolucion IS NULL)', None),
    (6, 3, 2, '<--', 'Asignacion activa (no devuelta previamente)', None),
    (7, 2, 3, '->', '@transaction.atomic: UPDATE asignacion SET fecha_devolucion=NOW', 'atomic'),
    (8, 2, 3, '->', '@transaction.atomic: UPDATE activo SET estado=DISPONIBLE', 'atomic'),
    (9, 3, 2, '<--', 'COMMIT exitoso', 'atomic'),
    (10, 2, 1, '<--', '200 OK {asignacion actualizada}', None),
    (11, 1, 1, 'self', 'Actualizar tabla + DB.addLog("Devolucion")', None),
    (12, 1, 0, '<--', 'Notificacion: "Devolucion registrada"', None),
], ['Excepcion: si ya fue devuelta -> 400 "La devolucion ya fue registrada para esta asignacion"'])

# SD-06: Crear Sede
add_seq_diagram('SD-06: Crear Sede', 'CU-05', PL, [
    (1, 0, 1, '->', 'Completa formulario (nombre, ciudad, direccion, telefono). Clic "Guardar"', None),
    (2, 1, 1, 'self', 'Validar campos obligatorios (nombre, ciudad, direccion)', None),
    (3, 1, 2, '->', 'DB.createSede({nombre, ciudad, direccion, telefono})', None),
    (4, 2, 2, 'self', 'Generar ID autoincremental + guardar en localStorage', None),
    (5, 2, 1, '<--', 'Sede creada {id, nombre, ciudad, direccion, telefono}', None),
    (6, 1, 1, 'self', 'DB.addLog("Sede creada", nombre) + setRefresh()', None),
    (7, 1, 0, '<--', 'Notificacion: "Sede guardada exitosamente"', None),
], ['Nota: Solo usuarios con rol ADMIN pueden acceder a esta funcionalidad.',
    'Flujo alterno eliminar: si tiene dependencias -> "No se puede eliminar: tiene activos/empleados/departamentos"'])

# SD-07: Generar Reportes
add_seq_diagram('SD-07: Generar Reportes', 'CU-06', PU, [
    (1, 0, 1, '->', 'Accede al modulo "Reportes" desde el menu lateral', None),
    (2, 1, 2, '->', 'DB.getActivos() + filtrar por sede si es OPERADOR', None),
    (3, 2, 1, '<--', 'activos[] clasificados por estado (disponibles, asignados, danados)', None),
    (4, 1, 2, '->', 'DB.getAsignaciones() + DB.getEmpleados()', None),
    (5, 2, 1, '<--', 'asignaciones[] + empleados[]', None),
    (6, 1, 1, 'self', 'Calcular estadisticas: valor total, conteos por estado', None),
    (7, 1, 0, '<--', 'Renderizar: 4 tarjetas resumen + tabs (Disponibles/Asignados/Reparacion)', None),
    (8, 0, 1, '->', '[Opcional] Clic en "Exportar CSV"', None),
    (9, 1, 1, 'self', 'Generar Blob CSV con codificacion UTF-8 + BOM', None),
    (10, 1, 0, '<--', 'Descarga automatica: activos_reporte.csv', None),
])

# SD-08: Crear Usuario
add_seq_diagram('SD-08: Crear Usuario', 'CU-07', PL, [
    (1, 0, 1, '->', 'Completa formulario (usuario, email, contrasena, nombre, apellido, rol, sede)', None),
    (2, 1, 1, 'self', 'Validar: campos obligatorios, contrasena >= 6 caracteres', None),
    (3, 1, 2, '->', 'DB.getUserByName(username) - verificar unicidad', None),
    (4, 2, 1, '<--', 'null (username disponible)', None),
    (5, 1, 2, '->', 'DB.createUser({username, email, password, rol, sede_id})', None),
    (6, 2, 2, 'self', 'Generar ID + almacenar en localStorage', None),
    (7, 2, 1, '<--', 'Usuario creado {id, username, rol, sede_id}', None),
    (8, 1, 1, 'self', 'DB.addLog("Usuario creado", username) + setRefresh()', None),
    (9, 1, 0, '<--', 'Notificacion: "Usuario creado: [username]"', None),
], ['Nota: Solo usuarios con rol ADMIN pueden crear otros usuarios.',
    'Flujo alterno: username duplicado -> "Usuario ya existe"'])

output = 'documentacion/Diagramas_Secuencia_SIGAE.docx'
doc.save(output)
print(f'Documento generado: {output}')
print(f'Tablas: {len(doc.tables)}')
