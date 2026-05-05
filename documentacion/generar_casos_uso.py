"""Genera las tablas de casos de uso en formato Word exacto de la plantilla."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54)
    s.right_margin = Cm(2.54)

def set_cell(cell, text, bold=False, size=10, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    run.bold = bold
    if bg:
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        sh = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): bg,
            qn('w:val'): 'clear'
        })
        shading.append(sh)

def add_cu_table(data):
    """
    data = {
        'nombre': str,
        'version': str,
        'referencia': str,
        'autor': str,
        'fecha': str,
        'actores': str,
        'referencias': str,
        'precondiciones': str,
        'postcondicion': str,
        'flujo': [(paso, actor, desc), ...],
        'alternos': str,
        'excepciones': str,
    }
    """
    # Contar filas: encabezado(1) + version/ref(1) + autor/fecha(1) + actores(1) + refs(1)
    # + precond(1) + postcond(1) + flujo_header(1) + flujo_cols(1) + pasos(N) + alternos(1) + excepciones(1)
    n_pasos = len(data['flujo'])
    n_rows = 11 + n_pasos  # fijas + pasos

    table = doc.add_table(rows=n_rows, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Ajustar anchos
    for row in table.rows:
        for i, w in enumerate([Cm(3.5), Cm(6), Cm(3.5), Cm(4.5)]):
            row.cells[i].width = w

    row_idx = 0

    # Fila 0: Caso de Uso | Nombre
    a, b, c, d = table.rows[row_idx].cells
    a.merge(b)
    c.merge(d)
    set_cell(a, 'Caso de Uso', bold=True, bg='003087')
    a.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    set_cell(c, data['nombre'], bold=True)
    row_idx += 1

    # Fila 1: Version | V1 | Referencia | CU-XX
    a, b, c, d = table.rows[row_idx].cells
    set_cell(a, 'Versión', bold=True, bg='E8EEF5')
    set_cell(b, data['version'])
    set_cell(c, 'Referencia', bold=True, bg='E8EEF5')
    set_cell(d, data['referencia'])
    row_idx += 1

    # Fila 2: Autor | nombre | Fecha | fecha
    a, b, c, d = table.rows[row_idx].cells
    set_cell(a, 'Autor', bold=True, bg='E8EEF5')
    set_cell(b, data['autor'])
    set_cell(c, 'Fecha', bold=True, bg='E8EEF5')
    set_cell(d, data['fecha'])
    row_idx += 1

    # Fila 3: Actores (merge 4 cols -> 1+3)
    a, b, c, d = table.rows[row_idx].cells
    set_cell(a, 'Actores', bold=True, bg='E8EEF5')
    b.merge(c).merge(d)
    set_cell(b, data['actores'])
    row_idx += 1

    # Fila 4: Referencias
    a, b, c, d = table.rows[row_idx].cells
    set_cell(a, 'Referencias', bold=True, bg='E8EEF5')
    b.merge(c).merge(d)
    set_cell(b, data['referencias'])
    row_idx += 1

    # Fila 5: Precondiciones
    a, b, c, d = table.rows[row_idx].cells
    set_cell(a, 'Precondiciones', bold=True, bg='E8EEF5')
    b.merge(c).merge(d)
    set_cell(b, data['precondiciones'])
    row_idx += 1

    # Fila 6: Postcondicion
    a, b, c, d = table.rows[row_idx].cells
    set_cell(a, 'Postcondición', bold=True, bg='E8EEF5')
    b.merge(c).merge(d)
    set_cell(b, data['postcondicion'])
    row_idx += 1

    # Fila 7: Flujo Normal (header, merge all)
    a, b, c, d = table.rows[row_idx].cells
    a.merge(b).merge(c).merge(d)
    set_cell(a, 'Flujo Normal', bold=True, bg='003087')
    a.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    row_idx += 1

    # Fila 8: Paso | Actor | Descripcion (headers)
    a, b, c, d = table.rows[row_idx].cells
    set_cell(a, 'Paso', bold=True, bg='E8EEF5')
    set_cell(b, 'Actor', bold=True, bg='E8EEF5')
    c.merge(d)
    set_cell(c, 'Descripción', bold=True, bg='E8EEF5')
    row_idx += 1

    # Filas de pasos
    for paso, actor, desc in data['flujo']:
        a, b, c, d = table.rows[row_idx].cells
        set_cell(a, str(paso))
        a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell(b, actor)
        c.merge(d)
        set_cell(c, desc)
        row_idx += 1

    # Flujos alternos
    a, b, c, d = table.rows[row_idx].cells
    set_cell(a, 'Flujos alternos', bold=True, bg='E8EEF5')
    b.merge(c).merge(d)
    set_cell(b, data['alternos'])
    row_idx += 1

    # Excepciones
    a, b, c, d = table.rows[row_idx].cells
    set_cell(a, 'Excepciones', bold=True, bg='E8EEF5')
    b.merge(c).merge(d)
    set_cell(b, data['excepciones'])

    doc.add_paragraph()  # Espacio entre tablas


# ================================================================
# CU-01: Iniciar Sesion
# ================================================================
add_cu_table({
    'nombre': 'Iniciar Sesión',
    'version': 'V1',
    'referencia': 'CU-01',
    'autor': 'Cristian Ramírez Romero',
    'fecha': '15/04/2026',
    'actores': 'Administrador, Operador',
    'referencias': 'RF01',
    'precondiciones': 'El usuario debe tener una cuenta registrada en el sistema con username y contraseña válidos.',
    'postcondicion': 'El usuario accede al sistema con su rol correspondiente. Se generan tokens JWT (access y refresh) y el sidebar se adapta según el rol.',
    'flujo': [
        (1, 'Usuario', 'Accede a la pantalla de login del sistema.'),
        (2, 'Usuario', 'Ingresa su nombre de usuario y contraseña en los campos del formulario.'),
        (3, 'Usuario', 'Hace clic en el botón "Ingresar".'),
        (4, 'Sistema', 'Envía las credenciales al backend mediante POST /api/token/.'),
        (5, 'Sistema', 'Verifica el username y valida la contraseña contra el hash PBKDF2-SHA256 almacenado.'),
        (6, 'Sistema', 'Genera tokens JWT (access con expiración corta, refresh con expiración larga).'),
        (7, 'Sistema', 'Almacena los tokens en localStorage y carga la sesión en AuthContext.'),
        (8, 'Sistema', 'Redirige al Dashboard. Si es ADMIN muestra menú completo; si es OPERADOR oculta Administración.'),
    ],
    'alternos': 'Si las credenciales son inválidas, el sistema muestra "Usuario o contraseña incorrectos" sin revelar cuál de los dos es incorrecto. El usuario puede intentar nuevamente.',
    'excepciones': 'Si el servidor backend no está disponible, se muestra un mensaje genérico de error de conexión sin exponer detalles técnicos.',
})

# ================================================================
# CU-02: Gestionar Activos
# ================================================================
add_cu_table({
    'nombre': 'Gestionar Activos',
    'version': 'V1',
    'referencia': 'CU-02',
    'autor': 'Cristian Ramírez Romero',
    'fecha': '15/04/2026',
    'actores': 'Administrador, Operador',
    'referencias': 'RF02, RF03, RF04, RF05, RF06, RF18',
    'precondiciones': 'El usuario debe haber iniciado sesión exitosamente. Si es OPERADOR, el sistema filtrará automáticamente mostrando únicamente los activos de su sede asignada.',
    'postcondicion': 'El activo queda registrado, actualizado o eliminado en la base de datos. El sistema genera una entrada en el log de actividad con fecha, usuario y detalle de la operación.',
    'flujo': [
        (1, 'Usuario', 'Selecciona la opción "Activos" en el menú lateral del sistema.'),
        (2, 'Sistema', 'Consulta la base de datos y muestra la lista de activos. Si es OPERADOR filtra por sede_id. Columnas: nombre, tipo, serial, valor, sede, departamento, estado.'),
        (3, 'Usuario', 'Hace clic en el botón "+ Nuevo Activo".'),
        (4, 'Sistema', 'Despliega formulario inline: nombre, descripción, tipo, serial, valor, fecha de compra, sede y departamento.'),
        (5, 'Usuario', 'Completa todos los campos requeridos y hace clic en "Guardar".'),
        (6, 'Sistema', 'Ejecuta validaciones: serial único, valor mayor a cero, campos obligatorios completos.'),
        (7, 'Sistema', 'Crea el registro con estado DISPONIBLE por defecto. Registra la operación en el log de actividad.'),
        (8, 'Sistema', 'Muestra notificación de éxito y actualiza la tabla de activos.'),
    ],
    'alternos': 'Serial duplicado: muestra "Serial ya existe". Valor ≤ 0: muestra "Valor debe ser mayor a 0". Campos vacíos: muestra "Complete todos los campos". Edición: carga datos en formulario, mismas validaciones excluyendo propio registro. Eliminación: si tiene asignaciones activas muestra "No se puede eliminar: está asignado"; si no, solicita confirmación.',
    'excepciones': 'Sesión JWT expirada: redirige automáticamente al login. Error de base de datos: muestra mensaje genérico sin exponer detalles internos.',
})

# ================================================================
# CU-03: Gestionar Empleados
# ================================================================
add_cu_table({
    'nombre': 'Gestionar Empleados',
    'version': 'V1',
    'referencia': 'CU-03',
    'autor': 'Cristian Ramírez Romero',
    'fecha': '15/04/2026',
    'actores': 'Administrador, Operador',
    'referencias': 'RF07, RF08',
    'precondiciones': 'El usuario debe haber iniciado sesión. Si es OPERADOR, solo verá empleados de su sede asignada.',
    'postcondicion': 'El empleado queda registrado, actualizado, dado de baja o eliminado. La operación se registra en el log de actividad.',
    'flujo': [
        (1, 'Usuario', 'Selecciona "Empleados" en el menú lateral.'),
        (2, 'Sistema', 'Muestra lista de empleados con columnas: nombre, cédula, sede, departamento, cargo, activos asignados, fecha ingreso y estado.'),
        (3, 'Usuario', 'Hace clic en "+ Nuevo Empleado".'),
        (4, 'Sistema', 'Despliega formulario: nombre, apellido, cédula, email, sede, departamento y cargo.'),
        (5, 'Usuario', 'Completa los campos obligatorios y hace clic en "Guardar".'),
        (6, 'Sistema', 'Valida que la cédula no esté duplicada y que los campos obligatorios estén completos.'),
        (7, 'Sistema', 'Crea el registro con estado activo=true y fecha_ingreso automática. Registra en log.'),
        (8, 'Sistema', 'Muestra notificación de éxito y actualiza la tabla.'),
    ],
    'alternos': 'Cédula duplicada: muestra "Cédula ya existe". Dar de baja: registra fecha_salida y marca activo=false. Reactivar: limpia fecha_salida y marca activo=true. Eliminar: si tiene activos asignados muestra "No se puede eliminar: tiene activos asignados"; si no, solicita confirmación.',
    'excepciones': 'Sesión JWT expirada: redirige al login.',
})

# ================================================================
# CU-04: Asignar Activo a Empleado
# ================================================================
add_cu_table({
    'nombre': 'Asignar Activo a Empleado',
    'version': 'V1',
    'referencia': 'CU-04',
    'autor': 'Cristian Ramírez Romero',
    'fecha': '15/04/2026',
    'actores': 'Administrador, Operador',
    'referencias': 'RF09, RF10, RF17, RF20',
    'precondiciones': 'Debe existir al menos un activo con estado DISPONIBLE y al menos un empleado registrado. Para OPERADOR, ambos deben pertenecer a su sede asignada.',
    'postcondicion': 'Para asignación: el activo cambia a estado ASIGNADO y se crea registro con fecha automática. Para devolución: el activo vuelve a DISPONIBLE y se registra la fecha de devolución.',
    'flujo': [
        (1, 'Usuario', 'Accede al módulo "Asignaciones" desde el menú lateral.'),
        (2, 'Sistema', 'Muestra formulario de nueva asignación (activo, empleado, observaciones) y tabla de historial ordenada por fecha descendente.'),
        (3, 'Usuario', 'Selecciona un activo del dropdown (solo muestra activos DISPONIBLES).'),
        (4, 'Usuario', 'Selecciona un empleado del dropdown.'),
        (5, 'Usuario', 'Opcionalmente escribe una observación y hace clic en "Registrar".'),
        (6, 'Sistema', 'Valida que el activo seleccionado siga en estado DISPONIBLE.'),
        (7, 'Sistema', 'Ejecuta transacción atómica: crea registro de asignación (fecha_asignacion=ahora) y actualiza estado del activo a ASIGNADO.'),
        (8, 'Sistema', 'Registra en log de actividad y muestra notificación de éxito.'),
        (9, 'Usuario', '[Devolución] Localiza asignación activa en la tabla y hace clic en "Devolver".'),
        (10, 'Sistema', 'Muestra diálogo de confirmación.'),
        (11, 'Usuario', 'Confirma la devolución.'),
        (12, 'Sistema', 'Ejecuta transacción atómica: registra fecha_devolucion=ahora y cambia estado del activo a DISPONIBLE.'),
        (13, 'Sistema', 'Registra en log y muestra notificación de éxito.'),
    ],
    'alternos': 'Sin activos disponibles: dropdown muestra "Sin disponibles" y botón se deshabilita. Activo ya asignado por otro usuario: muestra "Activo no disponible".',
    'excepciones': 'Devolución duplicada: "La devolución ya fue registrada para esta asignación". Fallo en transacción atómica: ninguna operación se aplica, datos quedan consistentes.',
})

# ================================================================
# CU-05: Gestionar Sedes (Solo Admin)
# ================================================================
add_cu_table({
    'nombre': 'Gestionar Sedes',
    'version': 'V1',
    'referencia': 'CU-05',
    'autor': 'Cristian Ramírez Romero',
    'fecha': '15/04/2026',
    'actores': 'Administrador',
    'referencias': 'RF11, RF19',
    'precondiciones': 'El usuario debe tener rol ADMIN. Los operadores no ven esta opción en el menú ni pueden acceder a la funcionalidad.',
    'postcondicion': 'Sede creada, editada o eliminada. La operación queda registrada en el log de actividad.',
    'flujo': [
        (1, 'Admin', 'Selecciona "Sedes" en el menú lateral, sección Organización.'),
        (2, 'Sistema', 'Muestra tabla de sedes: nombre, ciudad, dirección, teléfono, cantidad de departamentos y activos.'),
        (3, 'Admin', 'Hace clic en "+ Nueva Sede".'),
        (4, 'Sistema', 'Despliega formulario: nombre, ciudad, dirección y teléfono.'),
        (5, 'Admin', 'Completa los campos obligatorios y hace clic en "Guardar".'),
        (6, 'Sistema', 'Valida campos obligatorios, crea la sede y registra en log de actividad.'),
        (7, 'Sistema', 'Muestra notificación de éxito y actualiza la tabla.'),
    ],
    'alternos': 'Eliminar: el sistema verifica dependencias. Si la sede tiene activos, empleados o departamentos asociados, muestra "No se puede eliminar: la sede tiene [activos/empleados/departamentos] asociados" y bloquea la operación. Solo permite eliminar sedes sin dependencias.',
    'excepciones': 'Sesión JWT expirada: redirige al login.',
})

# ================================================================
# CU-06: Generar Reportes
# ================================================================
add_cu_table({
    'nombre': 'Generar Reportes',
    'version': 'V1',
    'referencia': 'CU-06',
    'autor': 'Cristian Ramírez Romero',
    'fecha': '15/04/2026',
    'actores': 'Administrador, Operador',
    'referencias': 'RF13, RF14',
    'precondiciones': 'El usuario debe haber iniciado sesión. Los datos se filtran automáticamente por sede si es OPERADOR.',
    'postcondicion': 'El usuario visualiza los reportes solicitados. Opcionalmente descarga un archivo CSV con los datos.',
    'flujo': [
        (1, 'Usuario', 'Selecciona "Reportes" en el menú lateral.'),
        (2, 'Sistema', 'Muestra 4 tarjetas de resumen: valor total, total activos, asignados y en reparación.'),
        (3, 'Sistema', 'Muestra tabs de navegación: Disponibles, Asignados, Reparación.'),
        (4, 'Usuario', 'Selecciona un tab para ver el reporte correspondiente.'),
        (5, 'Sistema', 'Muestra tabla con activos filtrados por estado: nombre, tipo, serial, valor y sede.'),
        (6, 'Usuario', 'Opcionalmente hace clic en "Exportar CSV".'),
        (7, 'Sistema', 'Genera archivo CSV con codificación UTF-8 (BOM) y lo descarga automáticamente.'),
    ],
    'alternos': 'Si no hay datos para el filtro seleccionado, el sistema muestra "Sin datos" en la tabla correspondiente.',
    'excepciones': 'Sesión JWT expirada: redirige al login.',
})

# ================================================================
# CU-07: Gestionar Usuarios (Solo Admin)
# ================================================================
add_cu_table({
    'nombre': 'Gestionar Usuarios',
    'version': 'V1',
    'referencia': 'CU-07',
    'autor': 'Cristian Ramírez Romero',
    'fecha': '15/04/2026',
    'actores': 'Administrador',
    'referencias': 'RF15, RF16',
    'precondiciones': 'El usuario debe tener rol ADMIN. La sección Administración no es visible para operadores.',
    'postcondicion': 'Usuario creado, editado o eliminado con rol y sede asignados. La operación queda registrada en el log de actividad.',
    'flujo': [
        (1, 'Admin', 'Accede al módulo "Usuarios" desde la sección Administración del menú lateral.'),
        (2, 'Sistema', 'Muestra formulario de creación arriba y tabla de usuarios abajo: usuario, nombre, email, rol, sede, fecha.'),
        (3, 'Admin', 'Completa: usuario, email, contraseña, nombre, apellido, rol (Operador/Administrador) y sede.'),
        (4, 'Admin', 'Hace clic en "Crear Usuario".'),
        (5, 'Sistema', 'Valida: username único, contraseña mínimo 6 caracteres, email con formato válido.'),
        (6, 'Sistema', 'Crea el usuario con contraseña cifrada (PBKDF2-SHA256) y registra en log.'),
        (7, 'Sistema', 'Muestra notificación de éxito y actualiza la tabla de usuarios.'),
    ],
    'alternos': 'Username duplicado: muestra "Usuario ya existe". Contraseña corta: muestra "Contraseña mínimo 6 caracteres". Edición: carga datos, contraseña vacía = no cambiar. Eliminar propia cuenta: muestra "No puede eliminar su propia cuenta" y bloquea.',
    'excepciones': 'Sesión JWT expirada: redirige al login.',
})

# Guardar
output = 'documentacion/Casos_de_Uso_SIGAE.docx'
doc.save(output)
print(f'Documento generado: {output}')
print(f'Tablas: {len(doc.tables)}')
